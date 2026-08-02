from docx import Document
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import re
import os
import datetime

import streamlit as st
from dotenv import load_dotenv
import google.generativeai as genai
from gtts import gTTS
import plotly.express as px
import plotly.graph_objects as go
import pandas as pd


# -----------------------------
# Create PDF
# -----------------------------
def create_pdf(content):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = []
    for line in content.split("\n"):
        story.append(Paragraph(line if line.strip() else "&nbsp;", styles["BodyText"]))
    doc.build(story)
    buffer.seek(0)
    return buffer


# -----------------------------
# Create DOCX
# -----------------------------
def create_docx(content):
    document = Document()
    document.add_heading("AI YouTube Writer", level=1)
    for line in content.split("\n"):
        document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


# -----------------------------
# Create Audio (gTTS Text-to-Speech)
# -----------------------------
def create_audio(text, tld="com", slow=False):
    # Remove markdown titles and hashtags for smooth audio output
    clean_text = re.sub(r"[#*`_]", "", text)
    tts = gTTS(text=clean_text[:4000], lang="en", tld=tld, slow=slow)
    buffer = BytesIO()
    tts.write_to_fp(buffer)
    buffer.seek(0)
    return buffer


# -----------------------------
# API Setup
# -----------------------------
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")

if api_key:
    genai.configure(api_key=api_key)
    try:
        model = genai.GenerativeModel("gemini-2.5-flash")
    except Exception:
        model = genai.GenerativeModel("gemini-1.5-flash")
else:
    model = None


# -----------------------------
# Page Configuration
# -----------------------------
st.set_page_config(
    page_title="AI YouTube Creator Studio",
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded"
)


# -----------------------------
# Session State Initialization
# -----------------------------
if "history" not in st.session_state:
    st.session_state.history = []

if "temperature" not in st.session_state:
    st.session_state.temperature = 0.7

if "max_tokens" not in st.session_state:
    st.session_state.max_tokens = 2000

if "writing_style" not in st.session_state:
    st.session_state.writing_style = "Professional"

if "last_generated_text" not in st.session_state:
    st.session_state.last_generated_text = ""


# -----------------------------
# CSS Loader
# -----------------------------
def load_css():
    if os.path.exists("assets/style.css"):
        with open("assets/style.css") as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)


load_css()


# -----------------------------
# Sidebar Navigation
# -----------------------------
with st.sidebar:
    st.markdown(
        """
        <div style='text-align: center; margin-bottom: 15px;'>
            <span class='animated-badge'>⚡ AI CREATOR STUDIO</span>
            <h2 style='font-size: 1.7rem; margin-top: 10px; margin-bottom: 0;'>🎬 YouTube Writer</h2>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.divider()

    page = st.radio(
        "🚀 Navigation",
        [
            "🏠 Home",
            "📊 Animated Dashboard",
            "✍️ Generate Script",
            "🔊 Text to Audio",
            "📜 Generation History",
            "⚙️ Settings",
        ],
    )

    st.divider()

    st.markdown(
        """
        <div class='custom-card' style='padding: 12px; font-size: 0.85rem;'>
            <strong style='color: #38bdf8;'>💡 Pro Creator Tip:</strong><br>
            Use <strong>Text to Audio</strong> to render high quality AI voiceovers for YouTube Shorts & Videos!
        </div>
        """,
        unsafe_allow_html=True,
    )


# -----------------------------
# Home Page
# -----------------------------
if page == "🏠 Home":
    st.markdown(
        """
        <div style='background: linear-gradient(135deg, rgba(99,102,241,0.2) 0%, rgba(236,72,153,0.15) 100%); 
                    padding: 30px; border-radius: 20px; border: 1px solid rgba(255,255,255,0.1); margin-bottom: 25px;'>
            <span class='animated-badge'>🔥 PREMIUM AI PLATFORM</span>
            <h1 style='font-size: 2.8rem; margin-top: 10px;'>🎬 AI YouTube Studio & Voiceover Generator</h1>
            <p style='font-size: 1.15rem; color: #cbd5e1; max-width: 800px;'>
                Craft viral YouTube scripts, storytelling narratives, thumbnail hooks, and instant AI voiceovers in seconds powered by Google Gemini AI & gTTS.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("📝 Scripts Generated", f"{len(st.session_state.history)}")

    with col2:
        total_words = sum(
            len(re.findall(r"\w+", item["content"]))
            for item in st.session_state.history
        )
        st.metric("📚 Total Words", f"{total_words:,}")

    with col3:
        st.metric("⚡ Engine", "Gemini 2.5 Flash")

    with col4:
        st.metric("🔊 Audio Engine", "gTTS Studio")

    st.divider()

    st.markdown("### 🌟 Key Studio Features")

    f_col1, f_col2, f_col3 = st.columns(3)

    with f_col1:
        st.markdown(
            """
            <div class='custom-card'>
                <h3>🎯 Viral Hook & Script Engine</h3>
                <p style='color: #94a3b8;'>Generates high-retention titles, description, SEO tags, storytelling outlines, and line-by-line scripts.</p>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with f_col2:
        st.markdown(
            """
            <div class='custom-card'>
                <h3>📊 Interactive Animated Analytics</h3>
                <p style='color: #94a3b8;'>Track script production, word counts, tone distribution, and channel growth projection with dynamic charts.</p>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with f_col3:
        st.markdown(
            """
            <div class='custom-card'>
                <h3>🔊 Text-to-Audio Synthesizer</h3>
                <p style='color: #94a3b8;'>Convert scripts to realistic natural human audio with multiple global accents (US, UK, IN, AU, CA).</p>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.divider()



# -----------------------------
# Animated Dashboard Page
# -----------------------------
elif page == "📊 Animated Dashboard":
    st.markdown(
        """
        <div style='margin-bottom: 20px;'>
            <span class='animated-badge'>📊 DYNAMIC METRICS</span>
            <h1 style='margin-top: 5px;'>Interactive Analytics Dashboard</h1>
            <p style='color: #94a3b8;'>Visual insights into your AI script production, content balance, and predicted growth.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Real & Simulated Metrics
    total_scripts = len(st.session_state.history)
    total_words = sum(
        len(re.findall(r"\w+", item["content"])) for item in st.session_state.history
    )
    est_audio_mins = round(total_words / 150, 1)  # average 150 words per minute
    avg_script_len = (
        int(total_words / total_scripts) if total_scripts > 0 else 0
    )

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("🎬 Total Scripts", f"{total_scripts}")
    m2.metric("📝 Total Words Created", f"{total_words:,}")
    m3.metric("🎙️ Est. Voiceover Duration", f"{est_audio_mins} mins")
    m4.metric("📊 Avg Script Length", f"{avg_script_len} words")

    st.divider()

    dash_col1, dash_col2 = st.columns(2)

    with dash_col1:
        st.subheader("📈 Script Production & Word Growth")

        # Generate sample timeline data based on session history or mock progression
        dates = [
            datetime.date.today() - datetime.timedelta(days=i)
            for i in range(6, -1, -1)
        ]
        if total_scripts == 0:
            words_series = [150, 300, 750, 1200, 1800, 2400, 3100]
            scripts_series = [1, 2, 4, 6, 8, 10, 13]
        else:
            cum_w = 0
            words_series = []
            scripts_series = []
            for i in range(7):
                cum_w += int(total_words / 7) + (i * 40)
                words_series.append(cum_w)
                scripts_series.append(min(total_scripts, i + 1))

        df_growth = pd.DataFrame(
            {
                "Date": [d.strftime("%b %d") for d in dates],
                "Cumulative Words": words_series,
                "Scripts Published": scripts_series,
            }
        )

        fig_growth = px.area(
            df_growth,
            x="Date",
            y="Cumulative Words",
            text="Cumulative Words",
            title="Word Count Accumulation",
            color_discrete_sequence=["#6366f1"],
        )
        fig_growth.update_traces(
            mode="lines+markers", line=dict(width=3, shape="spline")
        )
        fig_growth.update_layout(
            template="plotly_dark",
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
            margin=dict(l=20, r=20, t=40, b=20),
        )
        st.plotly_chart(fig_growth, use_container_width=True)

    with dash_col2:
        st.subheader("🎭 Content Tone & Genre Balance")

        if total_scripts > 0:
            tones = [item["tone"] for item in st.session_state.history]
            df_tones = pd.DataFrame({"Tone": tones}).value_counts().reset_index()
            df_tones.columns = ["Tone", "Count"]
        else:
            df_tones = pd.DataFrame(
                {
                    "Tone": [
                        "Professional",
                        "Storytelling",
                        "Funny",
                        "Motivational",
                        "Casual",
                    ],
                    "Count": [4, 3, 2, 2, 1],
                }
            )

        fig_pie = px.pie(
            df_tones,
            names="Tone",
            values="Count",
            hole=0.4,
            title="Distribution by Content Tone",
            color_discrete_sequence=px.colors.sequential.RdBu,
        )
        fig_pie.update_traces(textposition="inside", textinfo="percent+label")
        fig_pie.update_layout(
            template="plotly_dark",
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
            margin=dict(l=20, r=20, t=40, b=20),
        )
        st.plotly_chart(fig_pie, use_container_width=True)

    st.divider()

    dash_col3, dash_col4 = st.columns(2)

    with dash_col3:
        st.subheader("🔥 AI Quality & Virality Meter")

        fig_gauge = go.Figure(
            go.Indicator(
                mode="gauge+number+delta",
                value=94.5,
                delta={"reference": 85, "increasing": {"color": "#10b981"}},
                gauge={
                    "axis": {"range": [0, 100], "tickwidth": 1},
                    "bar": {"color": "#ec4899"},
                    "bgcolor": "rgba(30,41,59,0.5)",
                    "bordercolor": "gray",
                    "steps": [
                        {"range": [0, 50], "color": "#ef4444"},
                        {"range": [50, 80], "color": "#f59e0b"},
                        {"range": [80, 100], "color": "#10b981"},
                    ],
                    "threshold": {
                        "line": {"color": "white", "width": 4},
                        "thickness": 0.75,
                        "value": 94.5,
                    },
                },
                title={"text": "Estimated Audience Retention Score (%)"},
            )
        )
        fig_gauge.update_layout(
            template="plotly_dark",
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
            height=300,
            margin=dict(l=20, r=20, t=50, b=20),
        )
        st.plotly_chart(fig_gauge, use_container_width=True)

    with dash_col4:
        st.subheader("🚀 Channel Growth Simulation")

        weeks = [f"Week {i}" for i in range(1, 9)]
        views_projected = [1200, 3500, 8900, 15400, 29000, 48000, 76000, 125000]
        subs_projected = [45, 120, 310, 680, 1250, 2100, 3400, 5800]

        df_sim = pd.DataFrame(
            {
                "Week": weeks,
                "Projected Views": views_projected,
                "Projected Subscribers": subs_projected,
            }
        )

        fig_sim = px.bar(
            df_sim,
            x="Week",
            y="Projected Views",
            hover_data=["Projected Subscribers"],
            title="Estimated Views per Week with Regular AI Uploads",
            color="Projected Views",
            color_continuous_scale="Viridis",
        )
        fig_sim.update_layout(
            template="plotly_dark",
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
            height=300,
            margin=dict(l=20, r=20, t=50, b=20),
        )
        st.plotly_chart(fig_sim, use_container_width=True)


# -----------------------------
# Generate Script Page
# -----------------------------
elif page == "✍️ Generate Script":
    st.markdown(
        """
        <div style='margin-bottom: 20px;'>
            <span class='animated-badge'>✨ GEMINI AI ENGINE</span>
            <h1 style='margin-top: 5px;'>✍️ AI Script & Story Generator</h1>
            <p style='color: #94a3b8;'>Fill in the details below to produce structured YouTube scripts with SEO hooks.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.form("script_generator_form"):
        topic = st.text_input(
            "📝 Topic / Keyword",
            placeholder="Example: How Artificial Intelligence will change the world by 2030",
        )

        c1, c2, c3 = st.columns(3)

        with c1:
            tone = st.selectbox(
                "🎭 Select Tone",
                [
                    "Professional",
                    "Casual",
                    "Funny",
                    "Emotional",
                    "Motivational",
                    "Storytelling",
                    "Suspenseful",
                ],
            )

        with c2:
            genre = st.selectbox(
                "📚 Select Genre",
                [
                    "Technology",
                    "Education",
                    "Entertainment",
                    "Finance & Crypto",
                    "Travel & Vlogs",
                    "Gaming",
                    "Motivation & Self-Help",
                    "Storytelling",
                ],
            )

        with c3:
            content_type = st.radio(
                "🎬 Content Format",
                ["YouTube Script", "Story"],
                horizontal=True,
            )

        submit_btn = st.form_submit_button(
            "🚀 Generate Script & Content", use_container_width=True
        )

    if submit_btn:
        if not topic.strip():
            st.warning("⚠️ Please enter a topic before generating.")
        elif not model:
            st.error(
                "❌ GOOGLE_API_KEY is not configured in your .env file! Please add it in Settings or .env."
            )
        else:
            prompt = f"""
You are an expert viral YouTube Content Creator and Screenwriter.

Create a high-retention {content_type}.

Topic: {topic}
Tone: {tone}
Genre: {genre}
Writing Style: {st.session_state.writing_style}
Creativity Index: {st.session_state.temperature}
Target Max Length: {st.session_state.max_tokens} words

Formatting Structure Required:

# 🎯 SEO Title Options (Give 3 catchy options)

# 🖼 Thumbnail Text Ideas (Give 3 text overlay ideas)

# 📝 Optimized Description & Timestamps

# 🏷 High-Ranking SEO Tags

# 🔥 Hook (First 15 Seconds Speech Script)

# 📋 Content Outline

# 📖 Complete Script / Story (Detailed speech lines with sound/visual directions like [Visual: ...] or [Sound Effect: ...])

Make the script engaging, direct, and structured for maximum viewer retention.
"""
            try:
                with st.spinner("🤖 Gemini AI is crafting your masterpiece..."):
                    response = model.generate_content(prompt)
                    output = response.text

                st.session_state.last_generated_text = output
                st.session_state.history.append(
                    {
                        "topic": topic,
                        "tone": tone,
                        "genre": genre,
                        "type": content_type,
                        "content": output,
                        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                    }
                )

                st.success("✅ Generated Successfully!")
                st.divider()

                st.markdown(output)

                word_count = len(re.findall(r"\w+", output))
                st.info(
                    f"📊 Word Count: **{word_count} words** (~{round(word_count/150, 1)} minutes audio duration)"
                )

                pdf_file = create_pdf(output)
                docx_file = create_docx(output)

                st.markdown("### 📥 Download & Synthesize")

                d1, d2, d3, d4 = st.columns(4)

                with d1:
                    st.download_button(
                        "📄 Download PDF",
                        data=pdf_file,
                        file_name=f"script_{topic[:15]}.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                    )

                with d2:
                    st.download_button(
                        "📝 Download DOCX",
                        data=docx_file,
                        file_name=f"script_{topic[:15]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

                with d3:
                    st.download_button(
                        "📥 Download TXT",
                        data=output,
                        file_name=f"script_{topic[:15]}.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )

                with d4:
                    if st.button(
                        "🔊 Convert to Audio Now", use_container_width=True
                    ):
                        st.info("Head over to the 🔊 **Text to Audio** page to customize and download voiceovers!")

            except Exception as e:
                st.error(f"❌ Generation Error: {e}")


# -----------------------------
# Text to Audio Page
# -----------------------------
elif page == "🔊 Text to Audio":
    st.markdown(
        """
        <div style='margin-bottom: 20px;'>
            <span class='animated-badge'>🔊 AUDIO STUDIO</span>
            <h1 style='margin-top: 5px;'>Text-to-Audio Voiceover Synthesizer</h1>
            <p style='color: #94a3b8;'>Convert script lines or custom text into clear human AI voiceovers using gTTS.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Pre-fill with last generated text if available
    default_text = st.session_state.last_generated_text if st.session_state.last_generated_text else "Welcome to our channel! In today's video, we are exploring the revolutionary future of Artificial Intelligence and technology."

    audio_text = st.text_area(
        "📝 Script Text for Audio Synthesis",
        value=default_text,
        height=220,
        help="Paste script lines here to generate natural MP3 audio speech.",
    )

    c_a1, c_a2, c_a3 = st.columns(3)

    with c_a1:
        accent = st.selectbox(
            "🌍 Voice Accent / Region",
            [
                "🇺🇸 United States (com)",
                "🇬🇧 United Kingdom (co.uk)",
                "🇮🇳 India (co.in)",
                "🇦🇺 Australia (com.au)",
                "🇨🇦 Canada (ca)",
            ],
        )
        tld_map = {
            "🇺🇸 United States (com)": "com",
            "🇬🇧 United Kingdom (co.uk)": "co.uk",
            "🇮🇳 India (co.in)": "co.in",
            "🇦🇺 Australia (com.au)": "com.au",
            "🇨🇦 Canada (ca)": "ca",
        }
        selected_tld = tld_map[accent]

    with c_a2:
        speed = st.selectbox("⚡ Narration Speed", ["Normal", "Slow Speed"])
        is_slow = True if speed == "Slow Speed" else False

    with c_a3:
        st.write("")
        st.write("")
        synth_btn = st.button("🔊 Synthesize Audio (MP3)", use_container_width=True)

    if synth_btn:
        if not audio_text.strip():
            st.warning("⚠️ Please provide text to convert to audio.")
        else:
            try:
                with st.spinner("🎙️ Synthesizing voiceover audio..."):
                    audio_buffer = create_audio(
                        audio_text, tld=selected_tld, slow=is_slow
                    )

                st.success("✅ Voiceover Audio Synthesized Successfully!")

                # Audio wave animation display
                st.markdown(
                    """
                    <div class='sound-wave'>
                        <div class='sound-bar'></div>
                        <div class='sound-bar'></div>
                        <div class='sound-bar'></div>
                        <div class='sound-bar'></div>
                        <div class='sound-bar'></div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                st.audio(audio_buffer, format="audio/mp3")

                st.download_button(
                    "📥 Download Voiceover MP3",
                    data=audio_buffer,
                    file_name="youtube_voiceover.mp3",
                    mime="audio/mp3",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"❌ Audio Synthesis Error: {e}")


# -----------------------------
# Generation History Page
# -----------------------------
elif page == "📜 Generation History":
    st.markdown(
        """
        <div style='margin-bottom: 20px;'>
            <span class='animated-badge'>📜 CREATOR ARCHIVE</span>
            <h1 style='margin-top: 5px;'>Script Generation History</h1>
            <p style='color: #94a3b8;'>Review and re-export previously generated content.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if len(st.session_state.history) == 0:
        st.info("ℹ️ No content generated yet in this session. Head over to ✍️ **Generate Script** to begin!")
    else:
        if st.button("🗑 Clear History"):
            st.session_state.history = []
            st.success("History Cleared!")
            st.rerun()

        st.divider()

        for idx, item in enumerate(reversed(st.session_state.history), start=1):
            with st.expander(
                f"🎬 #{idx} | {item['topic']} ({item.get('date', 'Recent')})"
            ):
                col_h1, col_h2, col_h3 = st.columns(3)
                col_h1.write(f"**Tone:** {item['tone']}")
                col_h2.write(f"**Genre:** {item['genre']}")
                col_h3.write(f"**Format:** {item['type']}")

                st.divider()
                st.markdown(item["content"])

                # Direct download options
                h_pdf = create_pdf(item["content"])
                h_docx = create_docx(item["content"])

                hd1, hd2, hd3 = st.columns(3)
                hd1.download_button(
                    "📄 Download PDF",
                    data=h_pdf,
                    file_name=f"history_{idx}.pdf",
                    mime="application/pdf",
                    key=f"pdf_{idx}",
                )
                hd2.download_button(
                    "📝 Download DOCX",
                    data=h_docx,
                    file_name=f"history_{idx}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"docx_{idx}",
                )
                hd3.download_button(
                    "📥 Download TXT",
                    data=item["content"],
                    file_name=f"history_{idx}.txt",
                    mime="text/plain",
                    key=f"txt_{idx}",
                )


# -----------------------------
# Settings Page
# -----------------------------
elif page == "⚙️ Settings":
    st.markdown(
        """
        <div style='margin-bottom: 20px;'>
            <span class='animated-badge'>⚙️ SYSTEM CONFIG</span>
            <h1 style='margin-top: 5px;'>AI Settings</h1>
            <p style='color: #94a3b8;'>Adjust generation parameters and preferred writing style.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.subheader("🎨 Model Generation Parameters")

    c_s1, c_s2 = st.columns(2)

    with c_s1:
        temperature = st.slider(
            "🎨 Creativity (Temperature)",
            0.0,
            1.0,
            st.session_state.temperature,
            0.1,
            help="Higher values make output more creative and randomized.",
        )
        st.session_state.temperature = temperature

    with c_s2:
        max_tokens = st.slider(
            "📄 Max Output Length (Tokens)",
            500,
            5000,
            st.session_state.max_tokens,
            100,
        )
        st.session_state.max_tokens = max_tokens

    styles = ["Simple", "Professional", "Creative", "Storytelling", "Persuasive"]
    writing_style = st.selectbox(
        "✍️ Preferred Writing Style",
        styles,
        index=styles.index(st.session_state.writing_style),
    )
    st.session_state.writing_style = writing_style

    st.divider()

    st.success("✅ Settings Saved for this Session")
    